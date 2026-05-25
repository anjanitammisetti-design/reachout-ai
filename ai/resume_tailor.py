# ai/resume_tailor.py

import os
import uuid
import json
from datetime import datetime, timezone
from pathlib import Path
from google.cloud import bigquery
import vertexai
from vertexai.generative_models import GenerativeModel
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env")

if not os.getenv("GCP_PROJECT_ID"):
    os.environ["GCP_PROJECT_ID"] = "reachout-ai-496806"
    os.environ["GCP_REGION"]     = "us-central1"
    os.environ["BQ_DATASET"]     = "reachout"

PROJECT  = os.getenv("GCP_PROJECT_ID")
REGION   = os.getenv("GCP_REGION")
DATASET  = os.getenv("BQ_DATASET")
TABLE    = f"{PROJECT}.{DATASET}.tailored_resumes"

vertexai.init(project=PROJECT, location=REGION)
model     = GenerativeModel("gemini-2.5-flash")
bq_client = bigquery.Client(project=PROJECT)


# ── Helpers ───────────────────────────────────────────────────────────────────

def load_resume_text(path: str = "my_resume.txt") -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def fetch_jd(job_id: str) -> dict:
    query = f"""
        SELECT title, company, description_raw
        FROM `{PROJECT}.{DATASET}.raw_jobs`
        WHERE job_id = '{job_id}'
        LIMIT 1
    """
    rows = list(bq_client.query(query).result())
    if not rows:
        raise Exception(f"Job ID not found: {job_id}")
    return dict(rows[0])


# ── Gemini — tailor content ───────────────────────────────────────────────────

def get_tailored_content(resume_text: str, jd_text: str,
                          job_title: str, company: str) -> dict:
    prompt = f"""
You are an expert resume writer helping a Senior Data Engineer tailor their resume.

ORIGINAL RESUME:
{resume_text}

TARGET JOB TITLE: {job_title}
TARGET COMPANY: {company}

JOB DESCRIPTION:
{jd_text[:3000]}

Rewrite ONLY these sections to match the job description keywords.
Keep ALL facts TRUE — do not invent experience.
Write in clear, professional, human-readable English.
Quantify achievements where possible.

Return ONLY valid JSON — no markdown, no explanation:
{{
  "subtitle": "Senior Data Engineer | GCP | BigQuery | <add 1-2 relevant keywords from JD>",
  "summary": "Paragraph 1 — overall experience and expertise tailored to this role.\n\nParagraph 2 — specific technical strengths relevant to this JD.\n\nParagraph 3 — one confident sentence about career break: I took a planned break to relocate from India to Sydney, navigate the Australian PR process, and welcome my first child. I am now actively returning to data engineering and have spent this time building real skills in Vertex AI, dbt, and AI-powered pipelines.",
  "skills": {{
    "Cloud & Data Platform": ["skill1", "skill2"],
    "Data Engineering": ["skill1", "skill2"],
    "Data Warehousing & Modeling": ["skill1", "skill2"],
    "Data Processing & Orchestration": ["skill1", "skill2"],
    "Programming": ["skill1", "skill2"],
    "Governance & Data Quality": ["skill1", "skill2"],
    "DevOps & Tools": ["skill1", "skill2"],
    "Methodology & Collaboration": ["skill1", "skill2"]
  }},
  "experience": [
    {{
      "title": "Senior Data Engineer",
      "company": "Tata Consultancy Services – Bed Bath & Beyond",
      "duration": "Mar 2022 – Jan 2024",
      "bullets": [
        "Rewritten bullet 1 using JD keywords — keep facts true",
        "Rewritten bullet 2",
        "Rewritten bullet 3",
        "Rewritten bullet 4",
        "Rewritten bullet 5"
      ]
    }},
    {{
      "title": "PL/SQL Developer",
      "company": "Cognizant – Total Gas & Power",
      "duration": "Aug 2018 – Mar 2022",
      "bullets": [
        "Rewritten bullet 1 using JD keywords — keep facts true",
        "Rewritten bullet 2",
        "Rewritten bullet 3"
      ]
    }}
  ],
  "key_achievements": [
    "Achieved 95%+ performance improvement on critical workloads.",
    "Awarded On The Spot Award for successful delivery of project under challenging timelines."
  ],
  "keywords_added": ["keyword1", "keyword2", "keyword3"]
}}

WRITING STYLE RULES:
- Never use the word Seasoned — use Experienced or start with the job title directly
- Never use dynamic, passionate, results-driven, leverage, or synergy
- Write like a real person, not a LinkedIn cliche
- Keep the same confident tone as the original resume
- Start the summary with Senior Data Engineer with 5.5+ years of experience
- Each bullet point should be a clear complete sentence
- Quantify wherever the original resume has numbers
"""

    response = model.generate_content(prompt)
    text = response.text.strip()

    if "```" in text:
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]

    return json.loads(text.strip())


# ── Word document builder ─────────────────────────────────────────────────────

def build_docx(tailored: dict, job_title: str, company: str) -> str:
    """Build Word document matching Anjani's exact resume style."""

    doc = Document()
    # Set default document font to Calibri
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:asciiTheme'), 'Calibri')
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin   = Pt(54)
        section.right_margin  = Pt(54)

    # ── Helper: add a text run to a paragraph ─────────────
    def add_run(para, text, bold=False, size=10):
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        return run

    # ── Helper: add section heading with bottom border ─────
    def add_heading(title):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        # Black bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Helper: add bullet point ───────────────────────────
    def add_bullet(text, size=10):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        # Apply bullet numbering
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        return p

    # ── Name ──────────────────────────────────────────────
    name_p = doc.add_paragraph(style="Normal")
    name_p.paragraph_format.space_after = Pt(0)
    add_run(name_p, "TAMMISETTI ANJANI", bold=True, size=14)

    # ── Subtitle ───────────────────────────────────────────
    sub_p = doc.add_paragraph(style="Normal")
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(0)
    add_run(sub_p,
        tailored.get("subtitle",
            "Senior Data Engineer | GCP | BigQuery | Cloud Data Platforms | ETL/ELT Engineering"),
        bold=True, size=10)

    # ── Contact ────────────────────────────────────────────
    contact_p = doc.add_paragraph(style="Normal")
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after  = Pt(6)
    add_run(contact_p,
        "Sydney, NSW | Permanent Resident (SC 309)\n"
        "+61 422 480 664 | anjani.tammisetti@gmail.com\n"
        "LinkedIn: https://www.linkedin.com/in/anjani-tammisetti-8ab793114/\n"
        "GitHub: https://github.com/anjanitammisetti-design/reachout-ai",
        bold=False, size=10)

    # ── 1. Professional Summary ────────────────────────────
    add_heading("PROFESSIONAL SUMMARY")
    summary = tailored.get("summary", "")
    for para_text in summary.split("\n\n"):
        if para_text.strip():
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.space_after = Pt(4)
            add_run(p, para_text.strip(), bold=False, size=10)

    # ── 2. Technical Expertise ─────────────────────────────
    add_heading("TECHNICAL EXPERTISE")
    skills = tailored.get("skills", {})
    if isinstance(skills, dict):
        for category, skill_list in skills.items():
            sp = doc.add_paragraph(style="Normal")
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after  = Pt(2)
            add_run(sp, f"{category}: ", bold=True, size=10)
            add_run(sp, ", ".join(skill_list), bold=False, size=10)
    else:
        sp = doc.add_paragraph(style="Normal")
        add_run(sp, str(skills), bold=False, size=10)

    # ── 3. Professional Experience ─────────────────────────
    add_heading("PROFESSIONAL EXPERIENCE")

    # Career break — always first
    cb_title = doc.add_paragraph(style="Normal")
    cb_title.paragraph_format.space_before = Pt(6)
    cb_title.paragraph_format.space_after  = Pt(0)
    add_run(cb_title, "Career Break", bold=True, size=10)

    cb_company = doc.add_paragraph(style="Normal")
    cb_company.paragraph_format.space_before = Pt(0)
    cb_company.paragraph_format.space_after  = Pt(2)
    add_run(cb_company,
        "Relocation to Sydney, Australia | Australian Permanent Residency | Parental Leave\n"
        "Jan 2024 – Present",
        bold=False, size=10)

    for bullet in [
        "Relocated internationally from India to Sydney, Australia and successfully obtained Australian Permanent Residency (SC 309).",
        "Took planned parental leave following the birth of first child.",
        "Maintained technical currency through self-directed upskilling in Vertex AI, dbt, Gemini, and modern AI data stack.",
        "Currently building ReachOut AI — an end-to-end AI-powered job search tool using GCP, BigQuery, Airflow, dbt, and Gemini.",
    ]:
        add_bullet(bullet)

    # Other roles
    for exp in tailored.get("experience", []):
        tp = doc.add_paragraph(style="Normal")
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(0)
        add_run(tp, exp.get("title", ""), bold=True, size=10)

        cp = doc.add_paragraph(style="Normal")
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(2)
        add_run(cp,
            f"{exp.get('company','')}\n{exp.get('duration','')}",
            bold=False, size=10)

        for bullet in exp.get("bullets", []):
            add_bullet(bullet)

    # ── 4. Education ───────────────────────────────────────
    add_heading("EDUCATION")
    edu_p = doc.add_paragraph(style="Normal")
    edu_p.paragraph_format.space_after = Pt(0)
    add_run(edu_p,
        "Bachelor of Technology (B.Tech) – Electronics & Communication Engineering",
        bold=False, size=10)

    edu_d = doc.add_paragraph(style="Normal")
    edu_d.paragraph_format.space_before = Pt(0)
    add_run(edu_d,
        "Gudlavalleru Engineering College\n"
        "Affiliated to Jawaharlal Nehru Technological University\n"
        "2014 – 2018",
        bold=False, size=10)

    # ── 5. Certification ───────────────────────────────────
    add_heading("CERTIFICATION")
    cert_p = doc.add_paragraph(style="Normal")
    add_run(cert_p,
        "Google Associate Cloud Engineer (ACE)",
        bold=False, size=10)

    # ── 6. Key Achievements ────────────────────────────────
    add_heading("KEY ACHIEVEMENTS")
    for ach in tailored.get("key_achievements", [
        "Achieved 95%+ performance improvement on critical workloads.",
        "Awarded \"On The Spot Award\" for successful delivery of project under challenging timelines.",
    ]):
        add_bullet(ach)

    # ── Save file ──────────────────────────────────────────
    os.makedirs("outputs/resumes", exist_ok=True)
    company_short   = company[:20].strip().replace(' ', '_')
    job_title_short = job_title[:20].strip().replace(' ', '_')
    filename = f"outputs/resumes/resume_{company_short}_{job_title_short}.docx"
    doc.save(filename)
    return filename


# ── BigQuery save ─────────────────────────────────────────────────────────────

def save_to_bigquery(job_id: str, original: str,
                      tailored: dict, filename: str):
    row = {
        "resume_id":       str(uuid.uuid4()),
        "job_id":          job_id,
        "original_resume": original[:10000],
        "tailored_resume": json.dumps(tailored),
        "docx_filename":   filename,
        "created_at":      datetime.now(timezone.utc).isoformat(),
    }
    errors = bq_client.insert_rows_json(TABLE, [row])
    if errors:
        raise Exception(f"BigQuery error: {errors}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=== ReachOut AI — Resume Tailor ===\n")

    query = f"""
        SELECT job_id, title, company
        FROM `{PROJECT}.{DATASET}.raw_jobs`
        ORDER BY ingested_at DESC
        LIMIT 10
    """
    rows = list(bq_client.query(query).result())

    print("Your recent jobs:\n")
    for i, row in enumerate(rows, 1):
        print(f"  {i}. {row['company']} — {row['title']}")

    print()
    choice   = input("Enter the number of the job to tailor resume for: ").strip()
    selected = rows[int(choice) - 1]
    job_id   = selected["job_id"]

    print(f"\nTailoring for: {selected['company']} — {selected['title']}")
    print("Sending to Gemini... (10-20 seconds)\n")

    job      = fetch_jd(job_id)
    resume   = load_resume_text()
    tailored = get_tailored_content(
        resume,
        job["description_raw"],
        job["title"],
        job["company"]
    )

    print(f"Keywords added: {', '.join(tailored.get('keywords_added', []))}")

    filename = build_docx(tailored, job["title"], job["company"])
    print(f"\nWord document saved: {filename}")

    save_to_bigquery(job_id, resume, tailored, filename)
    print("Saved to BigQuery.")

    print(f"\nOpen your tailored resume:")
    print(f"  {os.path.abspath(filename)}")